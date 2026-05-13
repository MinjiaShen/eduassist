"""
post_processor.py — 模块 D：整合引擎（字段合并 + 模板渲染 + 导出）

提供三个核心函数：
- generate_case(): 用 Jinja2 模板渲染结构化医案
- export_docx():   将 Markdown 文本转为 DOCX 文档
- save_output():   保存输出文件到 output/ 目录
"""

from __future__ import annotations

import logging
import os
import re
from datetime import datetime
from pathlib import Path

import jinja2
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# 项目根目录（基于文件位置推断）
PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = PROJECT_ROOT / "output"
TEMPLATE_DIR = PROJECT_ROOT / "templates"

# ---------------------------------------------------------------------------
# 内部工具
# ---------------------------------------------------------------------------

# 需要拆分为列表的字段（inline_tag / highlight 类型）
# 包含单数和复数形式，因为 marker_parser 按 YAML field 名返回单数形式
_LIST_FIELD_NAMES = {
    "acupoint", "acupoints",
    "key_symptom", "key_symptoms",
    "treatment_direction", "treatment_directions",
    "unmatched_sections",
    "inline_tags",
    "highlights",
}

# 复诊字段前缀（用于自动组装 followup_visits 列表）
_FOLLOWUP_PREFIX = "followup_visit_"


def _build_followup_visits(fields: dict) -> list[dict]:
    """
    将 followup_visit_2, followup_visit_3 等独立字段
    组装为 followup_visits 列表，供模板循环渲染。
    """
    visits = []
    for key in sorted(fields.keys()):
        if key.startswith(_FOLLOWUP_PREFIX):
            num = key[len(_FOLLOWUP_PREFIX):]
            content = fields[key].strip()
            if content:
                visits.append({
                    "visit_label": f"{'初' if num == '1' else '二' if num == '2' else '三' if num == '3' else '四' if num == '4' else '五' if num == '5' else num}诊",
                    "changes": content,
                    "analysis": "",
                    "prescription": "",
                    "advice": "",
                })
    return visits


def _ensure_list_fields(fields: dict) -> dict:
    """
    处理 inline_tag / highlight 类型字段：
    - 如果值已是列表，直接保留
    - 如果是字符串，按换行 / 分号 / 逗号拆分为列表
    - 组装 followup_visits 列表
    """
    processed = dict(fields)
    for key, value in processed.items():
        if key in _LIST_FIELD_NAMES or key.endswith("_list"):
            if isinstance(value, str):
                # 按换行、分号、逗号拆分，过滤空串
                items = re.split(r"[;，,\n]+", value)
                processed[key] = [item.strip() for item in items if item.strip()]
            elif not isinstance(value, list):
                processed[key] = [str(value)] if value else []

    # 组装复诊列表
    if "followup_visits" not in processed:
        visits = _build_followup_visits(processed)
        if visits:
            processed["followup_visits"] = visits

    return processed


def _md_to_docx(doc: Document, md_text: str) -> None:
    """
    将 Markdown 文本基本映射到 python-docx Document 中。
    支持：# 标题、- 列表项、> 引用、普通段落、**粗体**、--- 分隔线。
    """
    lines = md_text.split("\n")
    in_list = False

    for line in lines:
        stripped = line.strip()

        # 空行 → 列表结束后添加段落间距
        if not stripped:
            if in_list:
                in_list = False
            continue

        # 分隔线
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 40)
            continue

        # 标题 (# ~ ####)
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            in_list = False
            continue

        # 列表项 (- / * / • / △ / ★ / →)
        list_match = re.match(r"^[-*•△★→]\s+(.*)", stripped)
        if list_match:
            text = list_match.group(1).strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, text)
            in_list = True
            continue

        # 引用块 (> ...)
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_runs(p, stripped)


def _add_runs(paragraph, text: str) -> None:
    """解析行内 Markdown 格式（**粗体**、*斜体*）并添加到段落。"""
    # 简单的 **bold** 和 *italic* 解析
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# 公开接口
# ---------------------------------------------------------------------------


def generate_case(fields: dict, template_path: str = "templates/case_template.md") -> str:
    """
    用 Jinja2 模板渲染结构化医案。

    Args:
        fields: 从 marker_parser.parse_text() 得到的 fields 字典
        template_path: 模板路径（相对于项目根目录或绝对路径）

    Returns:
        渲染后的 Markdown 文本；出错时返回包含错误信息的提示文本
    """
    try:
        # 解析模板路径
        tpl_path = Path(template_path)
        if not tpl_path.is_absolute():
            tpl_path = PROJECT_ROOT / tpl_path

        if not tpl_path.exists():
            err = f"模板文件不存在: {tpl_path}"
            logger.error(err)
            return f"<!-- 生成失败: {err} -->"

        # 创建 Jinja2 环境（从模板所在目录加载）
        env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(str(tpl_path.parent)),
            undefined=jinja2.Undefined,  # 未定义变量渲染为空
            keep_trailing_newline=True,
            trim_blocks=True,
            lstrip_blocks=True,
        )
        template = env.get_template(tpl_path.name)

        # 处理字段：拆分列表类型
        render_data = _ensure_list_fields(fields)

        # 添加时间戳
        render_data.setdefault(
            "generated_at",
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        )

        # 渲染
        result = template.render(**render_data)
        logger.info("医案模板渲染成功，字段数: %d", len(fields))
        return result

    except Exception as e:
        err = f"generate_case 异常: {e}"
        logger.exception(err)
        return f"<!-- 生成失败: {err} -->"


def export_docx(markdown_text: str, output_path: str) -> str:
    """
    将 Markdown 文本转换为 DOCX 文档并保存。

    Args:
        markdown_text: Markdown 格式文本
        output_path: 输出文件路径（绝对或相对项目根目录）

    Returns:
        输出文件的完整路径；出错时返回空字符串
    """
    try:
        out_path = Path(output_path)
        if not out_path.is_absolute():
            out_path = PROJECT_ROOT / out_path

        # 确保父目录存在
        out_path.parent.mkdir(parents=True, exist_ok=True)

        # 构建 DOCX
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(11)

        _md_to_docx(doc, markdown_text)

        doc.save(str(out_path))
        logger.info("DOCX 导出成功: %s", out_path)
        return str(out_path.resolve())

    except Exception as e:
        logger.exception("export_docx 异常: %s", e)
        return ""


def save_output(content: str, filename: str, fmt: str = "md") -> str:
    """
    保存内容到 output/ 目录，自动生成带时间戳的文件名。

    Args:
        content: 文件内容
        filename: 基础文件名（不含扩展名）
        fmt: 输出格式 ("md" / "docx" / "txt")

    Returns:
        输出文件的完整路径；出错时返回空字符串
    """
    try:
        # 确保 output 目录存在
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        # 生成带时间戳的文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r"[^\w\u4e00-\u9fff-]", "_", filename)
        full_name = f"{safe_name}_{timestamp}.{fmt}"
        out_path = OUTPUT_DIR / full_name

        if fmt == "docx":
            path = export_docx(content, str(out_path))
            if not path:
                return ""
            logger.info("文件已保存 (docx): %s", path)
            return path
        else:
            # md / txt 直接写文本
            out_path.write_text(content, encoding="utf-8")
            logger.info("文件已保存 (%s): %s", fmt, out_path)
            return str(out_path.resolve())

    except Exception as e:
        logger.exception("save_output 异常: %s", e)
        return ""


# ---------------------------------------------------------------------------
# 自测
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(name)s | %(levelname)s | %(message)s")

    # 测试字段
    test_fields = {
        "chief_complaint": "反复头痛3年，加重1周",
        "present_illness": "患者3年前开始出现头痛，以两侧太阳穴为主，近1周因工作压力大加重。",
        "past_history": "高血压病史5年",
        "tongue_diagnosis": "舌红，苔薄黄",
        "pulse_diagnosis": "脉弦细",
        "diagnosis": "头痛 — 肝阳上亢证",
        "treatment_principle": "平肝潜阳，熄风止痛",
        "prescription": "天麻钩藤饮加减",
        "acupoints": "百会; 太阳; 风池; 太冲",
        "key_symptoms": "头痛, 失眠, 烦躁",
        "treatment_directions": "平肝 → 潜阳\n止痛 → 安神",
        "doctor_advice": "忌辛辣，注意休息，保持情绪稳定",
        "unmatched_sections": "患者自述近期食欲尚可。",
    }

    print("=" * 60)
    print("测试 1: generate_case")
    print("=" * 60)
    md_result = generate_case(test_fields)
    print(md_result[:500], "..." if len(md_result) > 500 else "")

    print("\n" + "=" * 60)
    print("测试 2: save_output (md)")
    print("=" * 60)
    md_path = save_output(md_result, "测试医案", fmt="md")
    print(f"保存路径: {md_path}")

    print("\n" + "=" * 60)
    print("测试 3: save_output (docx)")
    print("=" * 60)
    docx_path = save_output(md_result, "测试医案", fmt="docx")
    print(f"保存路径: {docx_path}")

    print("\n" + "=" * 60)
    print("测试 4: save_output (txt)")
    print("=" * 60)
    txt_path = save_output(md_result, "测试医案", fmt="txt")
    print(f"保存路径: {txt_path}")

    print("\n✅ 所有测试完成")
